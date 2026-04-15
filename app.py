import warnings
import os
import re
import gc
import json
import io
import subprocess
from datetime import datetime, timedelta
import torch
import torchaudio

# DEFINITIVE FIX: Ensure AudioDecoder is accessible for Diarization
# Recent torchaudio versions (2.0+) renamed AudioDecoder to StreamReader
try:
    import torchaudio.io
    if hasattr(torchaudio.io, 'AudioDecoder'):
        from torchaudio.io import AudioDecoder
    elif hasattr(torchaudio.io, 'StreamReader'):
        from torchaudio.io import StreamReader as AudioDecoder
    else:
        AudioDecoder = None
except Exception:
    AudioDecoder = None

# Inject into builtins to satisfy libraries that access it globally
import builtins
builtins.AudioDecoder = AudioDecoder

# Increase HF timeout for better connectivity during model downloads
os.environ["HF_HUB_READ_TIMEOUT"] = "120"

# Suppress annoying environmental warnings
warnings.filterwarnings("ignore", category=UserWarning)
os.environ["PYTHONWARNINGS"] = "ignore"
os.environ["PYTORCH_CUDA_ALLOC_CONF"] = "expandable_segments:True"

from flask import Flask, request, jsonify, render_template, send_file
from faster_whisper import WhisperModel
from transformers import pipeline
from pyannote.audio import Pipeline
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt

# Ensure a stable backend
try:
    if hasattr(torchaudio, 'list_audio_backends') and "sox_io" in torchaudio.list_audio_backends():
        torchaudio.set_audio_backend("sox_io")
except Exception:
    pass

# Real-time job status tracking
# Format: { session_id: {"status": "...", "percent": 0} }
job_status = {}

def update_status(session_id, status, percent=None):
    """Updates job status and prints to log with percentage"""
    if session_id not in job_status:
        job_status[session_id] = {"status": "unknown", "percent": 0}
    
    if status is not None:
        job_status[session_id]["status"] = status
    if percent is not None:
        job_status[session_id]["percent"] = percent
        
    p_text = f" [{percent}%]" if percent is not None else ""
    print(f"🔄 [{session_id}] Status: {status}{p_text}")

# Load environment variables
load_dotenv()
HF_TOKEN = os.getenv("HF_TOKEN")

def clean_text_for_ai(text):
    """Remove timestamps [00:00], speaker labels [Speaker 1:], and extra whitespace"""
    if not text: return ""
    # Remove [00:00] or [00:00:00] or [0:00]
    text = re.sub(r'\[\d{1,2}:\d{2}(?::\d{2})?\]', '', text)
    # Remove Speaker labels [Speaker 1]: or Speaker 1:
    text = re.sub(r'\[?Speaker\s*\d+\]?:?', '', text)
    # Extra clean
    text = " ".join(text.split())
    return text

def chunk_text(text, max_words=500):
    """Splits text into chunks while preserving sentence integrity."""
    # Use punctuation as split points if possible
    sentences = text.replace('!', '.').replace('?', '.').split('.')
    chunks = []
    current_chunk = []
    current_word_count = 0
    
    for sentence in sentences:
        sentence = sentence.strip() + "."
        words = sentence.split()
        if not words: continue
        if current_word_count + len(words) > max_words and current_chunk:
            chunks.append(" ".join(current_chunk))
            current_chunk = words
            current_word_count = len(words)
        else:
            current_chunk.extend(words)
            current_word_count += len(words)
            
    if current_chunk:
        chunks.append(" ".join(current_chunk))
    return chunks

def generate_chapters(segments, summarizer):
    """Groups meeting segments into chapters with titles and timestamps"""
    if not segments:
        return []
    
    print("   [Intelligence] Generating chapters...")
    chapters = []
    current_block = []
    block_duration = 0
    block_start_time = segments[0]['start']

    # Group into roughly 2-minute blocks for chaptering
    for seg in segments:
        current_block.append(seg['text'])
        block_duration = seg['end'] - block_start_time
        
        if block_duration > 120 or seg == segments[-1]:
            text = " ".join(current_block)
            # PREMIUM HINGLISH PROMPT: Generate specific, catchy titles
            try:
                prompt = (
                    "Provide a very short, professional 3-4 word title for this meeting segment. "
                    "If it talks about money, mention Lakhs/Crores. If politics/decisions, use specific names. "
                    "Segment: " + text[:500]
                )
                res = summarizer(prompt, max_length=15, min_length=3, do_sample=False)
                title = res[0]['summary_text'].replace('Title:', '').replace('"', '').strip()
                # Clean up "The title is..." fluff
                title = re.sub(r'^(The title is|Title|Summary|Meeting|This segment is about)\s*:?\s*', '', title, flags=re.IGNORECASE)
            except:
                title = "Meeting Segment"
            
            chapters.append({
                "start": block_start_time,
                "title": title.title()
            })
            
            current_block = []
            block_start_time = seg['end']
            
    return chapters

def infer_speaker_roles(speaker_data):
    """Heuristically assigns roles to speakers based on participation patterns"""
    roles = {}
    if not speaker_data:
        return roles

    # Calculate stats per speaker
    speaker_stats = {}
    total_words = 0
    
    for speaker, text in speaker_data.items():
        words = text.split()
        word_count = len(words)
        total_words += word_count
        # Count questions
        questions = text.count('?')
        speaker_stats[speaker] = {"count": word_count, "questions": questions}

    # Sort by participation
    sorted_speakers = sorted(speaker_stats.items(), key=lambda x: x[1]['count'], reverse=True)
    
    if not sorted_speakers:
        return roles

    # Role 1: Lead Speaker (Most talkative)
    main_speaker = sorted_speakers[0][0]
    roles[main_speaker] = "Lead Speaker"
    
    # Role 2: The Inquirer (Most questions asked)
    max_q = -1
    inquirer = None
    for spk, stats in speaker_stats.items():
        if stats['questions'] > max_q:
            max_q = stats['questions']
            inquirer = spk
    if inquirer and max_q > 0:
        roles[inquirer] = "Active Inquirer" if inquirer != main_speaker else "Lead & Moderator"

    # Role 3: Participants
    for spk in speaker_data.keys():
        if spk not in roles:
            roles[spk] = "Contributor" if speaker_stats[spk]['count'] > (total_words * 0.1) else "Passive Participant"
            
    return roles

def recursive_summarize(text, summarizer, depth=0):
    """Recursively summarizes text, optimized for Indian business context."""
    words = text.split()
    
    # Base Case: Short enough for a final structured summary
    if len(words) < 600 or depth > 3:
        # PREMUM HINGLISH PROMPT: Preserve business terms and structure
        safe_input = (
            "Summarize this Indian business/political meeting accurately. Use professional English, "
            "but keep critical Hinglish terms if they add clarity. "
            "STRUCTURE: \n1. KEY DISCUSSION (Entities, Names, Geopolitics)\n"
            "2. NUMERICAL DATA (Targets, Budgets in Lakhs/Crores)\n"
            "3. FINAL DECISIONS & ACTION ITEMS\n\n"
            "Transcript:\n" + " ".join(words[:900])
        )
        try:
            res = summarizer(safe_input, max_length=400, min_length=150, do_sample=False)
            return res[0]["summary_text"]
        except Exception as e:
            print(f"Recursive base case failed: {e}")
            return "Transcript too complex for AI summary. Use the chapters and action HUD."

    # Recursive Pass (Map)
    print(f"   [Recursive Summary Depth {depth}] Mapping {len(words)} words...")
    chunks = chunk_text(text, max_words=450)
    summaries = []
    
    for i, chunk in enumerate(chunks):
        try:
            # Short summary for each chunk
            res = summarizer(f"Summarize briefly:\n{chunk}", max_length=100, min_length=30, do_sample=False)
            summaries.append(res[0]["summary_text"])
        except Exception:
            continue
            
    # Reduce Pass
    combined_summaries = " ".join(summaries)
    print(f"   [Recursive Summary Depth {depth}] Reducing results...")
    return recursive_summarize(combined_summaries, summarizer, depth + 1)

def extract_action_items(text):
    """Extract actionable tasks from transcript using regex patterns"""
    if not text:
        return []

    # Clean timestamps & speaker labels
    clean_text = clean_text_for_ai(text)

    # Split into sentences
    sentences = re.split(r'[.!?]\s+', clean_text)

    # Deep Hinglish commitment patterns
    action_patterns = [
        r"(?:we|i)\s+(?:will|need to|should|have to|must)\s+.+",
        r"(?:let's|please|kindly)\s+.+",
        r"(?:action item|next step|task)s?:?\s+.+",
        r"(?:assign|assigned to|responsibility of)\s+.+",
        r"(?:follow up on|follow-up on)\s+.+",
        r"(?:make sure to|ensure that|dhyan rakhein)\s+.+",
        r".+(?:karna padega|karna hai|manage karna)\b",
        r".+(?:targets?|deadline)\s+is\s+.+",
        r"decided to\s+.+"
    ]

    extracted = []

    for sentence in sentences:
        s = sentence.strip()
        if not s:
            continue

        s_lower = s.lower()

        for pattern in action_patterns:
            if re.search(pattern, s_lower):
                # Clean formatting
                cleaned = s.strip()

                # Capitalize properly
                cleaned = cleaned[0].upper() + cleaned[1:] if len(cleaned) > 1 else cleaned

                extracted.append(cleaned)
                break  # avoid duplicate matches

    # Remove duplicates while preserving order
    unique_items = list(dict.fromkeys(extracted))

    # Limit to top 10 (clean UX)
    return unique_items[:10]

def cleanup_old_media(days=7, execute=False):
    """Checks or deletes audio files older than specified days."""
    audio_dir = "recordings/audio"
    if not os.path.exists(audio_dir):
        return 0, 0
    
    now = datetime.now().timestamp()
    count = 0
    total_size = 0
    for f in os.listdir(audio_dir):
        file_path = os.path.join(audio_dir, f)
        if os.path.isfile(file_path):
            file_age = now - os.path.getmtime(file_path)
            if file_age > (days * 86400):
                size = os.path.getsize(file_path)
                total_size += size
                count += 1
                if execute:
                    try:
                        os.remove(file_path)
                    except Exception as e:
                        print(f"Cleanup error for {f}: {e}")
    
    return count, total_size

def cleanup_old_files():
    """Maintenance: Deletes recordings/summaries older than 24 hours to keep the system lean."""
    try:
        now = datetime.now()
        retention_limit = now - timedelta(hours=24)
        print("🔧 Running auto-cleanup utility (24h retention)...")
        
        directories = [
            "recordings/audio", 
            "recordings/transcript", 
            "recordings/summaries", 
            "recordings/metadata"
        ]
        
        count = 0
        for directory in directories:
            if not os.path.exists(directory): continue
            for filename in os.listdir(directory):
                file_path = os.path.join(directory, filename)
                if os.path.isfile(file_path):
                    file_time = datetime.fromtimestamp(os.path.getmtime(file_path))
                    if file_time < retention_limit:
                        os.remove(file_path)
                        count += 1
        
        if count > 0:
            print(f"✅ Cleanup complete: Removed {count} old meeting files.")
    except Exception as e:
        print(f"❌ Cleanup failed: {e}")

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 500 * 1024 * 1024  # 500MB limit
# Run maintenance on startup
cleanup_old_files()

# Device configuration
device = "cuda" if torch.cuda.is_available() else "cpu"

# Model Cache (Lazy Loading to save VRAM)
whisper_model_cache = {"name": None, "model": None}

def get_whisper_model(model_name):
    """Lazy loads faster-whisper models and clears cache to save VRAM"""
    global whisper_model_cache
    if whisper_model_cache["name"] == model_name:
        return whisper_model_cache["model"]
    
    # Unload previous model to free memory
    if whisper_model_cache["model"] is not None:
        del whisper_model_cache["model"]
        # CTranslate2 models aren't directly deleted like torch modules, 
        # but clearing the reference and calling gc helps
        import gc
        gc.collect()
        torch.cuda.empty_cache()
    
    print(f"Loading Faster-Whisper {model_name} model (Hinglish Intelligence)... ")
    try:
        # float16 is optimal for 4GB GPUs. int8_float16 is even faster if needed.
        compute_type = "float16" if device == "cuda" else "int8"
        model = WhisperModel(model_name, device=device, compute_type=compute_type)
        whisper_model_cache = {"name": model_name, "model": model}
        return model
    except Exception as e:
        print(f"CUDA load failed for {model_name}: {e}. Falling back to CPU (int8)... ")
        model = WhisperModel(model_name, device="cpu", compute_type="int8")
        whisper_model_cache = {"name": model_name, "model": model}
        return model

# Lazy Diarization Pipeline to save VRAM on 4GB cards
diarization_pipeline_cache = None

def get_diarization_pipeline():
    """Loads and moves Pyannote 3.1 to GPU only when needed"""
    global diarization_pipeline_cache
    if diarization_pipeline_cache:
        return diarization_pipeline_cache
    
    if not HF_TOKEN:
        return None

    print("Loading Diarization pipeline (shuffling to VRAM)...")
    try:
        pipeline_obj = Pipeline.from_pretrained(
            "pyannote/speaker-diarization-3.1",
            token=HF_TOKEN
        )
        # Always keep diarization on CPU (stable)
        pipeline_obj.to(torch.device("cpu"))
        diarization_pipeline_cache = pipeline_obj
        return pipeline_obj
    except Exception as e:
        print(f"Diarization load failed: {e}")
        return None

def unload_diarization():
    """Unloads diarization from VRAM"""
    global diarization_pipeline_cache
    if diarization_pipeline_cache:
        print("Unloading Diarization from VRAM to make room...")
        try:
            del diarization_pipeline_cache
        except:
            pass
        diarization_pipeline_cache = None
        torch.cuda.empty_cache()

# Lazy Load Summarizer to save VRAM on startup
summarizer_cache = None

def get_summarizer():
    """Lazy loads the summarization pipeline on demand"""
    global summarizer_cache
    if summarizer_cache:
        return summarizer_cache
    
    print("Loading Summarizer pipeline (on CPU to save VRAM)...")
    try:
        # Optimized for speed/memory: sshleifer/distilbart-cnn-12-6 is 40% smaller and 2x faster than large-cnn
        summarizer_cache = pipeline("summarization", model="sshleifer/distilbart-cnn-12-6", device=-1)
        return summarizer_cache
    except Exception as e:
        print(f"Warning: Summarizer load failed ({e}).")
        return None

def unload_summarizer():
    """Unloads summarizer from RAM if needed"""
    global summarizer_cache
    if summarizer_cache:
        del summarizer_cache
        summarizer_cache = None

import subprocess

def denoise_audio(input_path):
    """
    Apply Deep Learning Noise Cancellation (RNNNoise)
    Fixes Status 183 by converting WebM/Opus to WAV/PCM first.
    """
    # Generate stable intermediate paths
    base_path = os.path.splitext(input_path)[0]
    # Prevent duplicate suffixes
    if base_path.endswith("_pcm"): base_path = base_path[:-4]
    if base_path.endswith("_clean"): base_path = base_path[:-6]
    
    pcm_wav = f"{base_path}_pcm.wav"
    clean_wav = f"{base_path}_clean.wav"
    model_path = "models/rnnoise/bd.rnnn"
    
    print(f"Converting and cleaning audio: {input_path}")
    try:
        # Stage 1: Convert to high-fidelity PCM WAV (16kHz Mono) if not already in that format
        # We only run this if we're not already trying to overwrite our own input
        if os.path.abspath(input_path) != os.path.abspath(pcm_wav):
            subprocess.run([
                "ffmpeg", "-y", "-i", input_path,
                "-ar", "16000", "-ac", "1", "-c:a", "pcm_s16le", pcm_wav
            ], check=True, capture_output=True)
        else:
            pcm_wav = input_path # Already PCM

        # Stage 2: Apply RNNNoise if model exists
        if os.path.exists(model_path):
            try:
                model_abs_path = os.path.abspath(model_path)
                subprocess.run([
                    "ffmpeg", "-y", "-i", pcm_wav,
                    "-af", f"arnndn=m='{model_abs_path}'",
                    clean_wav
                ], check=True, capture_output=True)
                return clean_wav
            except Exception as denoise_err:
                print(f"Noise reduction failed (falling back to clean PCM): {denoise_err}")
                return pcm_wav
        
        return pcm_wav
    except Exception as e:
        print(f"Audio conversion failed: {e}")
        return input_path

def extract_head_tail(text, limit=400):
    """Extract first and last parts of a long meeting to capture Intro + Conclusion"""
    words = text.split()
    if len(words) <= limit * 2:
        return text
    
    head = " ".join(words[:limit])
    tail = " ".join(words[-limit:])
    return f"{head} ... [Long Meeting Continued] ... {tail}"

@app.route("/")
def index():
    return render_template("index.html")

import time

def format_timestamp(seconds):
    return time.strftime("%M:%S", time.gmtime(seconds))

def truncate_text(text, max_tokens=1024):
    # Rough estimate: 1 token ~= 4 characters or 0.75 words
    # We'll use a conservative limit of words to avoid needing the full tokenizer here
    words = text.split()
    if len(words) > 800: # Conservative limit to stay under 1024 tokens
        return " ".join(words[:800])
    return text

def generate_ai_title(text):
    """Generates a concise title using cleaned English text"""
    try:
        # We prefer the English version for titles
        clean_prompt = clean_text_for_ai(text)
        if not clean_prompt.strip():
            return "Empty Meeting"
        
        # Safe truncation for the model's window
        safe_text = truncate_text(clean_prompt, max_tokens=500)
        
        current_summarizer = get_summarizer()
        if not current_summarizer:
            return "New AI Meeting"
            
        res = current_summarizer(safe_text, max_length=15, min_length=10, do_sample=False)
        title = res[0]["summary_text"].strip().replace(".", "").title()
        
        # Limit to 5 words
        return " ".join(title.split()[:5])
    except:
        return "New AI Meeting"

def create_docx(session_id, title, content, type_name):
    """Helper to generate professional DOCX files"""
    doc = Document()
    doc.add_heading(title, 0)
    
    p = doc.add_paragraph()
    run = p.add_run(f"Meeting Report: {type_name.title()}")
    run.bold = True
    
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_heading(type_name.title(), level=1)
    
    # Add content
    doc.add_paragraph(content)
    
    # Save to buffer
    target_path = f"recordings/{'transcript' if 'transcript' in type_name.lower() else 'summaries'}/{session_id}.docx"
    doc.save(target_path)
    return target_path

@app.route("/process", methods=["POST"])
def process_audio():
    file = request.files["audio"]
    custom_title = request.form.get("title", "").strip()
    
    # UI Toggles from Settings
    model_quality = request.form.get("model_quality", "small").strip()
    if not model_quality or model_quality == "null":
        model_quality = "small"
        
    use_hinglish = request.form.get("hinglish", "true").lower() == "true"
    
    # "Turbo" speed logic: skip denoising if using the fastest model
    use_turbo = (model_quality == "base")
    
    # Use the ID passed from frontend for real-time tracking
    session_id = request.form.get("session_id", datetime.now().strftime("%Y%m%d%H%M%S"))
    
    # Create directories if they don't exist
    for d in ["recordings/audio", "recordings/transcript", "recordings/summaries", "recordings/metadata"]:
        os.makedirs(d, exist_ok=True)

    # Save original file with its extension
    orig_ext = os.path.splitext(file.filename)[1].lower()
    if not orig_ext: orig_ext = ".webm" # Default for browser recordings
    
    orig_filename = f"{session_id}_original{orig_ext}"
    orig_path = os.path.join("recordings/audio", orig_filename)
    file.save(orig_path)
    
    # Init progress
    update_status(session_id, "extracting", 5)
    print(f"--- Processing Session: {session_id} ({orig_ext}) Quality: {model_quality} ---")

    # Final PCM Audio Path for AI pipeline
    pcm_path = os.path.join("recordings/audio", f"{session_id}_pcm.wav")

    # Universal Media Extraction/Conversion
    video_extensions = ['.mp4', '.mkv', '.mov', '.avi', '.wmv', '.flv', '.webm']
    
    try:
        if orig_ext in video_extensions or orig_ext != ".wav":
            print(f"1. Extracting/Converting audio from {orig_ext}...")
            # Extract high-fidelity audio (16kHz, mono, pcm_s16le)
            subprocess.run([
                "ffmpeg", "-y", "-i", orig_path,
                "-vn", "-acodec", "pcm_s16le", "-ar", "16000", "-ac", "1", pcm_path
            ], check=True, capture_output=True)
            update_status(session_id, "extracting", 15)
            audio_path = pcm_path
        else:
            audio_path = orig_path

        # Deep Learning Noise Cancellation
        if not use_turbo:
            update_status(session_id, "denoising", 20)
            model_path = os.path.abspath("models/rnnoise/bd.rnnn")
            
            # Input to denoising should be the PCM version if it exists
            input_audio = pcm_path if os.path.exists(pcm_path) else audio_path
            clean_path = f"recordings/audio/session_{session_id.replace('session_', '')}_clean.wav"
            
            try:
                print("1a. Attempting RNN Noise Reduction...")
                # Filter syntax: ensure no spaces and absolute path
                # We use FFmpeg 7.x compliant syntax
                subprocess.run([
                    "ffmpeg", "-y", "-i", input_audio,
                    "-af", f"arnndn=model={model_path}",
                    clean_path
                ], check=True, capture_output=True)
                audio_path = clean_path
                update_status(session_id, "denoising", 30)
                print("RNN Noise Reduction Successful.")
            except Exception as e:
                print(f"RNN Noise Reduction failed (trying FFT fallback): {e}")
                try:
                    print("1b. Starting FFT Noise Smoothing (afftdn)...")
                    # afftdn is a high-quality, non-RNN fallback
                    subprocess.run([
                        "ffmpeg", "-y", "-i", input_audio,
                        "-af", "afftdn=nr=12:nf=-30",
                        clean_path
                    ], check=True, capture_output=True)
                    audio_path = clean_path
                    print("FFT Fallback Successful.")
                except Exception as ef:
                    print(f"All denoising failed (using PCM): {ef}")
                    audio_path = pcm_path if os.path.exists(pcm_path) else orig_path
        else:
            print("Turbo speed active: Skipping Noise Reduction")
            audio_path = pcm_path if os.path.exists(pcm_path) else orig_path
    except Exception as e:
        print(f"Extraction/Conversion failed: {e}")
        audio_path = orig_path
    
    # Validation: Ensure non-zero size
    if not os.path.exists(audio_path) or os.path.getsize(audio_path) == 0:
        print(f"!! Critical Error: Audio file '{audio_path}' is missing or empty.")
        job_status[session_id] = "failed"
        return jsonify({"error": "Failed to process audio. The file might be corrupted or empty."}), 400

    # Transcription with Dynamic Model Switching and Lazy Loading
    selected_model_name = model_quality
    
    # Pre-transcription memory cleanup: Ensure Diarization and previous models are purged
    unload_diarization()
    torch.cuda.empty_cache()

    update_status(session_id, "transcribing", 35)
    # Load only what we need
    whisper_model = get_whisper_model(selected_model_name)
    print(f"2. Starting Whisper Transcription ({selected_model_name} model)...")
    
    # Check for cancellation before heavy AI starts
    if job_status.get(session_id) == "cancelled":
        print(f"!! Job {session_id} cancelled before transcription.")
        return jsonify({"status": "cancelled"}), 200

    # Deep Prompt for Hinglish Accuracy
    current_prompt = (
        "Namaste, this is an Indian meeting. हम आज targets aur progress par discuss karenge. "
        "Please transcribe in Hinglish, capturing both English and Hindi words clearly. "
        "Avoid mapping Indian accents to rare African or Asian languages like Malagasy or Burmese."
    ) if use_hinglish else "This is an Indian English meeting with a standard accent. Please transcribe accurately."

    # Update progress
    update_status(session_id, "transcribing", 40)
    try:
        # Phase 1: Primary Transcription (Hinglish/Native)
        print("2a. Starting transcription pass...")
        try:
            # faster-whisper uses a segments generator for real-time progress!
            segments, info = whisper_model.transcribe(
                audio_path,
                initial_prompt=current_prompt,
                task="transcribe",
                beam_size=1,
                vad_filter=True,  # ⬇️ Ultra-effective noise filtering
                vad_parameters=dict(min_silence_duration_ms=500),
                temperature=0,
                condition_on_previous_text=False
            )
            
            detected_language = info.language
            total_duration = info.duration
            
            # Metadata extraction and hallucination guard
            hallucination_langs = ["mg", "my", "ps", "haw", "ln", "lo", "su", "mt", "so", "sn", "yo"]
            if detected_language in hallucination_langs:
                print(f"⚠️ Hallucination detected ({detected_language}). Falling back to en.")
                detected_language = "en"

            whisper_segments_orig = []
            print(f"--- Detected Language: {detected_language} (Final: {detected_language}) ---")
            
            # Pass 1: Transcribe segments with real-time feedback
            for segment in segments:
                # Update progress based on duration (Scale: 35% to 65%)
                if total_duration > 0:
                    intra_progress = 40 + (segment.start / total_duration) * 25
                    update_status(session_id, f"transcribing ({detected_language})", int(intra_progress))
                
                text = segment.text.strip()
                
                # REFINED HALLUCINATION GUARD: Filter out repetitive AI artifacts and nonsensical noise
                if text and len(set(text)) > 3 and len(text.split()) >= 2:
                    # Detect common Whisper hallucinations (repetitive small phrases)
                    if len(whisper_segments_orig) > 0:
                        last_text = whisper_segments_orig[-1]["text"]
                        if text == last_text: continue # Skip exact repeats
                        
                    # Format for internal compatibility
                    whisper_segments_orig.append({
                        "start": segment.start,
                        "end": segment.end,
                        "text": text
                    })
            
            raw_text = " ".join([s["text"] for s in whisper_segments_orig])

            # Phase 2: English Translation (Only if needed)
            whisper_segments_en = None
            if detected_language != "en":
                # Check for cancellation between passes
                if job_status.get(session_id) == "cancelled":
                    return jsonify({"status": "cancelled"}), 200
                
                print(f"2b. Performing translation to English with real-time feedback...")
                segments_en, info_en = whisper_model.transcribe(
                    audio_path,
                    initial_prompt="Transcribe accurately into English.",
                    task="translate",
                    beam_size=1 if selected_model_name in ["small", "base"] else 3,
                    vad_filter=True,
                    temperature=0
                )
                
                whisper_segments_en = []
                total_duration_en = info_en.duration
                
                for segment in segments_en:
                    # Update progress (Scale: 65% to 80%)
                    if total_duration_en > 0:
                        intra_progress = 65 + (segment.start / total_duration_en) * 15
                        update_status(session_id, "translating to en", int(intra_progress))
                    
                    t = segment.text.strip()
                    if t and len(set(t)) > 2 and len(t.split()) >= 2:
                        whisper_segments_en.append({
                            "start": segment.start,
                            "end": segment.end,
                            "text": t
                        })

        except Exception as e:
            print(f"Faster-Whisper error: {e}")
            raise e

    except Exception as e:
        print(f"Whisper transcription failed: {e}")
        job_status[session_id] = "failed"
        return jsonify({"error": f"Transcription failed: {str(e)}"}), 500

    # IMPORTANT: Free VRAM for the heavy Diarization pipeline
    # Unload heavy models (Small/Medium) to make room for Pyannote 3.1
    del whisper_model
    whisper_model_cache = {"name": None, "model": None}
    torch.cuda.empty_cache()

    # Update progress
    update_status(session_id, "diarizing", 75)
    print(f"3. Transcription complete. Length: {len(raw_text)}")

    import gc
    gc.collect()
    torch.cuda.empty_cache()
    
    # Check for cancellation before Diarization
    if job_status.get(session_id) == "cancelled":
        print(f"!! Job {session_id} cancelled before diarization.")
        torch.cuda.empty_cache()
        return jsonify({"status": "cancelled"}), 200
    
    # Validation: If no speech was detected, skip heavy processing
    if len(raw_text) < 5:
        print("!! No speech detected. Skipping Diarization/Summarization.")
        return jsonify({
            "id": session_id,
            "title": custom_title or "Short Recording",
            "transcript": "No speech detected in this recording.",
            "summary": "Audio was too short or silent to generate a summary."
        })

    # Release memory for Diarization
    if torch.cuda.is_available():
        torch.cuda.empty_cache()

    update_status(session_id, "diarizing", 85)

    # Diarization with Pyannote (Lazy load)
    transcript_original = ""
    transcript_en = ""
    current_pipeline = get_diarization_pipeline()
    if current_pipeline:
        print("4. Starting Diarization...")
        try:
            diarization = current_pipeline(audio_path)

            def label_segments(segments):
                output = ""
                # Performance Fix: iterate tracks once outside the loop
                diarization_turns = list(diarization.itertracks(yield_label=True))
                
                for segment in segments:
                    start, end, text = segment["start"], segment["end"], segment["text"].strip()
                    speakers_in_segment = []
                    
                    for turn, _, speaker in diarization_turns:
                        intersection = min(end, turn.end) - max(start, turn.start)
                        if intersection > 0:
                            speakers_in_segment.append((speaker, intersection))

                    if speakers_in_segment:
                        dominant_speaker = max(speakers_in_segment, key=lambda x: x[1])[0]
                        speaker_label = f"Person {dominant_speaker.split('_')[-1]}"
                    else:
                        speaker_label = "Unknown Speaker"

                    output += f"[{format_timestamp(start)}] {speaker_label}: {text}\n"
                return output

            transcript_original = label_segments(whisper_segments_orig)
            if whisper_segments_en:
                transcript_en = label_segments(whisper_segments_en)

            print("5. Diarization complete.")
            unload_diarization()

        except Exception as e:
            print(f"⚠️ Diarization failed: {e}")
            print("➡ Falling back to non-speaker transcript")

            transcript_original = "\n".join([
                f"[{format_timestamp(s['start'])}] {s['text'].strip()}"
                for s in whisper_segments_orig
            ])

            if whisper_segments_en:
                transcript_en = "\n".join([
                    f"[{format_timestamp(s['start'])}] {s['text'].strip()}"
                    for s in whisper_segments_en
                ])

            unload_diarization()
    else:
        transcript_original = "\n".join([f"[{format_timestamp(s['start'])}] {s['text'].strip()}" for s in whisper_segments_orig])
        if whisper_segments_en:
            transcript_en = "\n".join([f"[{format_timestamp(s['start'])}] {s['text'].strip()}" for s in whisper_segments_en])

    # Check for cancellation before Summarization
    if job_status.get(session_id) == "cancelled":
        print(f"!! Job {session_id} cancelled before summarization.")
        return jsonify({"status": "cancelled"}), 200

    # 5. Summary Generation & AI Insights
    # Clear VRAM before CPU heavy Summarization to ensure no OOM on 4GB cards
    torch.cuda.empty_cache()
    
    update_status(session_id, "summarizing", 90)
    summary = "Not enough speech detected to generate insights."
    chapters = []
    
    # Selection of Best Text: Prefer Full English Transcript without timestamps
    full_en_text = transcript_en or transcript_original
    clean_summary_input = clean_text_for_ai(full_en_text).strip()

    if len(clean_summary_input) >= 30:
        print("6. Generating High-Fidelity Chapters & Summary...")
        try:
            current_summarizer = get_summarizer()
            if current_summarizer:
                # Premium Chaptering (Intelligent Segments)
                chapters = generate_chapters(whisper_segments_en or whisper_segments_orig, current_summarizer)
                # Use robust recursive summarizer to handle any transcript length
                summary = recursive_summarize(clean_summary_input, current_summarizer)
        except Exception as e:
            print(f"Intelligence pass failed: {e}")
            summary = "Summary generation failed or was too complex."

    if not custom_title:
        # Use English cleaned text for the Title
        custom_title = generate_ai_title(full_en_text)

    # Speaker Role Profiling
    speaker_data = {}
    total_raw = transcript_en or transcript_original
    for line in total_raw.split('\n'):
        if ': ' in line:
            parts = line.split(': ')
            spk = parts[0].split('] ')[-1] if ']' in parts[0] else parts[0]
            txt = parts[1]
            speaker_data[spk] = speaker_data.get(spk, "") + " " + txt
    
    speaker_roles = infer_speaker_roles(speaker_data)

    # Unload summarizer to free RAM
    unload_summarizer()

    # Action Items Extraction from English Transcript
    action_items = extract_action_items(transcript_en or transcript_original)

    # Save Files
    with open(f"recordings/transcript/{session_id}.txt", "w") as f:
        f.write(transcript_original)
    if transcript_en:
        with open(f"recordings/transcript/{session_id}_en.txt", "w") as f:
            f.write(transcript_en)
            
    with open(f"recordings/summaries/{session_id}.txt", "w") as f:
        f.write(summary)
    
    metadata = {
        "id": session_id,
        "title": custom_title,
        "timestamp": datetime.now().strftime("%b %d, %Y %H:%M"),
        "raw_text_length": len(raw_text),
        "language": detected_language,
        "action_items": action_items,
        "chapters": chapters,
        "speaker_roles": speaker_roles
    }
    with open(f"recordings/metadata/{session_id}.json", "w") as f:
        json.dump(metadata, f)
    
    print(f"--- Finished Session: {session_id} ---\n")

    # Clean up status
    if session_id in job_status:
        update_status(session_id, "completed", 100)

    return jsonify({
        "id": session_id,
        "title": custom_title,
        "transcript": transcript_original,
        "transcript_en": transcript_en,
        "language": detected_language,
        "summary": summary,
        "action_items": action_items,
        "chapters": chapters,
        "speaker_roles": speaker_roles
    })

@app.route("/search")
def search_history():
    query = request.args.get("q", "").lower()
    results = []
    metadata_dir = "recordings/metadata"
    
    if os.path.exists(metadata_dir):
        files = sorted(os.listdir(metadata_dir), reverse=True)
        for f in files:
            if f.endswith(".json"):
                with open(os.path.join(metadata_dir, f), "r") as json_file:
                    data = json.load(json_file)
                    
                    # Search in Title
                    if query in data.get("title", "").lower():
                        results.append(data)
                        continue
                    
                    # Search in Transcript (optional/comprehensive)
                    transcript_path = f"recordings/transcript/{data['id']}.txt"
                    if os.path.exists(transcript_path):
                        with open(transcript_path, "r") as t_file:
                            if query in t_file.read().lower():
                                results.append(data)
                                
    return jsonify(results)

@app.route("/history", methods=["GET"])
def get_history():
    history = []
    metadata_dir = "recordings/metadata"
    if os.path.exists(metadata_dir):
        files = sorted(os.listdir(metadata_dir), reverse=True)
        for f in files:
            if f.endswith(".json"):
                with open(os.path.join(metadata_dir, f), "r") as json_file:
                    try:
                        history.append(json.load(json_file))
                    except:
                        continue
    return jsonify(history)

@app.route("/progress/<session_id>", methods=["GET"])
def get_progress(session_id):
    """Returns granular progress data (status + percent)."""
    data = job_status.get(session_id, {"status": "unknown", "percent": 0})
    return jsonify(data)

@app.route("/cancel/<session_id>", methods=["POST"])
def cancel_process(session_id):
    """Sets a cancellation flag for a background job."""
    print(f"🛑 Cancellation requested for: {session_id}")
    job_status[session_id] = "cancelled"
    return jsonify({"success": True})

@app.route("/admin/cleanup", methods=["GET", "POST"])
def admin_cleanup():
    """Manual cleanup handler."""
    if request.method == "POST":
        count, size = cleanup_old_media(execute=True)
        return jsonify({
            "success": True, 
            "message": f"Successfully deleted {count} files ({round(size / (1024*1024), 2)} MB cleared)."
        })
    else:
        # Just check status
        count, size = cleanup_old_media(execute=False)
        return jsonify({
            "count": count, 
            "size_mb": round(size / (1024*1024), 2)
        })

@app.route("/history/<session_id>", methods=["GET"])
def get_meeting(session_id):
    try:
        with open(f"recordings/transcript/{session_id}.txt", "r") as f:
            transcript = f.read()
        
        transcript_en = ""
        en_path = f"recordings/transcript/{session_id}_en.txt"
        if os.path.exists(en_path):
            with open(en_path, "r") as f:
                transcript_en = f.read()

        with open(f"recordings/summaries/{session_id}.txt", "r") as f:
            summary = f.read()
        
        metadata = {}
        if os.path.exists(f"recordings/metadata/{session_id}.json"):
            with open(f"recordings/metadata/{session_id}.json", "r") as f:
                metadata = json.load(f)

        return jsonify({
            "id": session_id,
            "title": metadata.get("title", session_id),
            "language": metadata.get("language", "en"),
            "transcript": transcript,
            "transcript_en": transcript_en,
            "summary": summary,
            "action_items": metadata.get("action_items", [])
        })
    except FileNotFoundError:
        return jsonify({"error": "Meeting not found"}), 404

@app.route("/history/<session_id>", methods=["PUT"])
def rename_meeting(session_id):
    new_title = request.json.get("title")
    meta_path = f"recordings/metadata/{session_id}.json"
    if os.path.exists(meta_path):
        with open(meta_path, "r") as f:
            data = json.load(f)
        data["title"] = new_title
        with open(meta_path, "w") as f:
            json.dump(data, f)
        return jsonify({"success": True})
    return jsonify({"error": "Not found"}), 404

@app.route("/history/<session_id>", methods=["DELETE"])
def delete_meeting(session_id):
    # Remove all related files (Original, English, Summaries, Audio versions)
    prefixes = [
        f"recordings/audio/{session_id}_original",
        f"recordings/audio/{session_id}_pcm",
        f"recordings/audio/{session_id}_clean",
        f"recordings/transcript/{session_id}",
        f"recordings/transcript/{session_id}_en",
        f"recordings/summaries/{session_id}",
        f"recordings/metadata/{session_id}"
    ]
    
    import glob
    for p in prefixes:
        for f in glob.glob(f"{p}.*"):
            try:
                os.remove(f)
            except:
                pass
    return jsonify({"success": True})

@app.route("/download/<format_type>/<type>/<session_id>")
def download_file(format_type, type, session_id):
    if format_type == "docx":
        # Generate on the fly to reflect latest edits
        try:
            with open(f"recordings/metadata/{session_id}.json", "r") as f:
                meta = json.load(f)
            
            content_path = f"recordings/{'transcript' if type == 'transcript' else 'summaries'}/{session_id}.txt"
            with open(content_path, "r") as f:
                content = f.read()
            
            docx_path = create_docx(session_id, meta['title'], content, type)
            return send_file(docx_path, as_attachment=True)
        except Exception as e:
            print(f"DOCX failed: {e}")
            return "Generation failed", 500
            
    # Standard Txt Download
    file_path = f"recordings/{'transcript' if type == 'transcript' else 'summaries'}/{session_id}.txt"
    if os.path.exists(file_path):
        return send_file(file_path, as_attachment=True)
    return "File not found", 404

@app.route("/recordings/audio/<session_id>")
def serve_audio(session_id):
    """Securely serve the processed audio for playback"""
    # Prefer cleaned version
    clean_path = f"recordings/audio/{session_id}_clean.wav"
    pcm_path = f"recordings/audio/{session_id}_pcm.wav"
    
    final_path = clean_path if os.path.exists(clean_path) else pcm_path
    
    if os.path.exists(final_path):
        return send_file(final_path, mimetype="audio/wav")
    return jsonify({"error": "Audio not found"}), 404

if __name__ == "__main__":
    cleanup_old_media()  # Run storage cleanup on startup
    app.run(debug=True, host="0.0.0.0")